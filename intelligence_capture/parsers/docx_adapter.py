"""
DOCX document adapter using python-docx

Parses Word documents and extracts text, headings, tables, and structure.
"""

from docx import Document
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

from .base_adapter import BaseAdapter
from ..models.document_payload import DocumentPayload


class DOCXAdapter(BaseAdapter):
    """
    Microsoft Word DOCX parsing adapter

    Uses python-docx to extract text, headings, paragraphs, and tables
    from .docx files. Preserves document structure and Spanish content.

    Example:
        >>> adapter = DOCXAdapter()
        >>> payload = adapter.parse(Path('procedimiento.docx'), metadata)
        >>> len(payload.tables)
        3
        >>> payload.sections[0]['title']
        'Procedimiento de Recepción'
    """

    @property
    def supported_mime_types(self) -> List[str]:
        """DOCX MIME types"""
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]

    def parse(
        self,
        file_path: Path,
        metadata: Dict[str, Any]
    ) -> DocumentPayload:
        """
        Parse DOCX document

        Extracts text, headings, tables, and preserves document structure.

        Args:
            file_path: Path to DOCX file
            metadata: Connector metadata

        Returns:
            DocumentPayload with extracted content

        Raises:
            ValueError: If DOCX parsing fails (Spanish error message)
        """
        # Validate inputs
        self.validate_metadata(metadata)

        if not file_path.exists():
            raise FileNotFoundError(
                f"Archivo DOCX no encontrado: {file_path}"
            )

        start_time = datetime.now()

        try:
            doc = Document(file_path)

            # Extract paragraphs and identify headings
            content_parts = []
            sections = []
            current_page = 1  # DOCX doesn't have page concept

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                content_parts.append(text)

                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    # Extract heading level from style name (e.g., "Heading 1" → 1)
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1

                    sections.append({
                        'title': text,
                        'level': level,
                        'page': current_page
                    })

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables, 1):
                # Extract headers from first row
                if len(table.rows) > 0:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                    # Extract data rows
                    rows = []
                    for row in table.rows[1:]:
                        rows.append([cell.text.strip() for cell in row.cells])

                    tables.append({
                        'headers': headers,
                        'rows': rows,
                        'page': current_page,
                        'table_index': table_idx
                    })

            # Combine content
            full_content = '\n\n'.join(content_parts)

            # Detect language
            language = self.detect_language(full_content)

            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()

            # Create payload
            return DocumentPayload(
                # Identity
                document_id=metadata['document_id'],
                org_id=metadata['org_id'],
                checksum=metadata['checksum'],

                # Source metadata
                source_type=metadata['source_type'],
                source_format='docx',
                mime_type=self.supported_mime_types[0],
                original_path=file_path,

                # Content (Spanish preserved)
                content=full_content,
                language=language,

                # Structure
                page_count=1,  # DOCX is continuous
                sections=sections,
                tables=tables,
                images=[],  # Image extraction requires specialized handling

                # Processing metadata
                context_tags=metadata.get('context_tags', []),
                processed_at=start_time,
                processing_time_seconds=processing_time,

                # Additional metadata
                metadata={
                    **metadata,
                    'parser': 'python-docx',
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                }
            )

        except Exception as e:
            raise ValueError(
                f"Error procesando DOCX {file_path.name}: {e}"
            )
